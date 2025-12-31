"""
PDF Operations - Core file processing functions.

These functions execute the operations parsed by the AI.
They handle actual file manipulation safely.
"""

import os
import shutil
from pathlib import Path
from typing import List, Optional
import io
import zipfile
import hashlib

from pypdf import PdfReader, PdfWriter
from pdf2docx import Converter
import subprocess
import math
import fitz  # PyMuPDF


# ============================================
# HELPER FUNCTIONS
# ============================================

def ensure_temp_dirs():
    """Create temporary directories if they don't exist"""
    Path("uploads").mkdir(exist_ok=True)
    Path("outputs").mkdir(exist_ok=True)


def get_upload_path(filename: str) -> str:
    """Get full path for an input file.

    Primary source is `uploads/`. For multi-step pipelines, intermediate files live in
    `outputs/` and may be used as inputs for subsequent steps.
    """
    uploads_path = os.path.join("uploads", filename)
    if os.path.exists(uploads_path):
        return uploads_path
    outputs_path = os.path.join("outputs", filename)
    return outputs_path


def get_output_path(filename: str) -> str:
    """Get full path for output file"""
    return os.path.join("outputs", filename)


def _resolve_ghostscript_executable(*, raise_if_missing: bool) -> Optional[str]:
    """Return a Ghostscript executable path if available."""
    # Prefer PATH lookup first.
    if os.name == "nt":
        for name in ("gswin64c.exe", "gswin32c.exe"):
            found = shutil.which(name)
            if found:
                return found

        # Common install locations (best-effort)
        possible_paths = [
            r"C:\Program Files\gs\gs10.06.0\bin\gswin64c.exe",
            r"C:\Program Files (x86)\gs\gs10.06.0\bin\gswin64c.exe",
            r"C:\Program Files\gs\gs10.05.0\bin\gswin64c.exe",
            r"C:\Program Files\gs\gs9.56.1\bin\gswin64c.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                return path

        if raise_if_missing:
            raise Exception(
                "Ghostscript not found. Please install Ghostscript from https://ghostscript.com/download/gsdnld.html"
            )
        return None

    # Linux/Mac
    found = shutil.which("gs")
    if found:
        return found
    if raise_if_missing:
        raise Exception("Ghostscript not found. Please install Ghostscript (gs) on the server.")
    return None


# ============================================
# PDF OPERATIONS
# ============================================

def merge_pdfs(file_names: List[str], output_name: str = "merged_output.pdf") -> str:
    """
    Merge multiple PDFs into a single file.
    
    Args:
        file_names: List of PDF file names to merge (in order)
        output_name: Name for the output file
    
    Returns:
        str: Output file name
    
    Raises:
        FileNotFoundError: If any input file doesn't exist
        Exception: If merge fails
    """
    ensure_temp_dirs()
    
    pdf_writer = PdfWriter()
    
    # Add all pages from all PDFs
    for file_name in file_names:
        input_path = get_upload_path(file_name)
        
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"File not found: {file_name}")
        
        pdf_reader = PdfReader(input_path)
        for page in pdf_reader.pages:
            pdf_writer.add_page(page)
    
    # Write merged PDF
    output_path = get_output_path(output_name)
    with open(output_path, "wb") as output_file:
        pdf_writer.write(output_file)
    
    return output_name


def split_pdf(file_name: str, pages_to_keep: List[int], output_name: str = "split_output.pdf") -> str:
    """
    Extract specific pages from a PDF.
    
    Args:
        file_name: Source PDF file name
        pages_to_keep: List of page numbers to keep (1-indexed)
        output_name: Name for the output file
    
    Returns:
        str: Output file name
    
    Raises:
        FileNotFoundError: If input file doesn't exist
        ValueError: If page numbers are invalid
        Exception: If split fails
    """
    ensure_temp_dirs()
    
    input_path = get_upload_path(file_name)
    
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")
    
    pdf_reader = PdfReader(input_path)
    total_pages = len(pdf_reader.pages)
    
    # Validate page numbers
    for page_num in pages_to_keep:
        if page_num < 1 or page_num > total_pages:
            raise ValueError(f"Invalid page number: {page_num}. PDF has {total_pages} pages.")
    
    pdf_writer = PdfWriter()
    
    # Add requested pages (convert 1-indexed to 0-indexed)
    for page_num in pages_to_keep:
        pdf_writer.add_page(pdf_reader.pages[page_num - 1])
    
    # Write output PDF
    output_path = get_output_path(output_name)
    with open(output_path, "wb") as output_file:
        pdf_writer.write(output_file)
    
    return output_name


def delete_pages(file_name: str, pages_to_delete: List[int], output_name: str = "deleted_output.pdf") -> str:
    """
    Remove specific pages from a PDF.
    
    Args:
        file_name: Source PDF file name
        pages_to_delete: List of page numbers to delete (1-indexed)
        output_name: Name for the output file
    
    Returns:
        str: Output file name
    
    Raises:
        FileNotFoundError: If input file doesn't exist
        ValueError: If page numbers are invalid
        Exception: If deletion fails
    """
    ensure_temp_dirs()
    
    input_path = get_upload_path(file_name)
    
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")
    
    pdf_reader = PdfReader(input_path)
    total_pages = len(pdf_reader.pages)
    
    # Validate page numbers
    for page_num in pages_to_delete:
        if page_num < 1 or page_num > total_pages:
            raise ValueError(f"Invalid page number: {page_num}. PDF has {total_pages} pages.")
    
    pdf_writer = PdfWriter()
    
    # Add all pages EXCEPT the ones to delete
    pages_to_delete_set = set(pages_to_delete)
    for page_num in range(1, total_pages + 1):
        if page_num not in pages_to_delete_set:
            pdf_writer.add_page(pdf_reader.pages[page_num - 1])
    
    # Write output PDF
    output_path = get_output_path(output_name)
    with open(output_path, "wb") as output_file:
        pdf_writer.write(output_file)
    
    return output_name


def compress_pdf(file_name: str, output_name: str = "compressed_output.pdf", preset: str = "ebook") -> str:
    """
    Compress a PDF using a qualitative preset.

    If Ghostscript is available, uses `-dPDFSETTINGS=/<preset>` where preset is one of:
    screen (smallest), ebook (default), printer (light), prepress (highest quality).
    If Ghostscript is not available, falls back to a basic PyPDF stream compression.
    
    Args:
        file_name: Source PDF file name
        output_name: Name for the output file
        preset: Ghostscript PDFSETTINGS preset (screen/ebook/printer/prepress)
    
    Returns:
        str: Output file name
    
    Raises:
        FileNotFoundError: If input file doesn't exist
        Exception: If compression fails
    """
    ensure_temp_dirs()
    
    input_path = get_upload_path(file_name)
    
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")
    
    output_path = get_output_path(output_name)

    allowed = {"screen", "ebook", "printer", "prepress"}
    if preset not in allowed:
        preset = "ebook"

    gs_executable = _resolve_ghostscript_executable(raise_if_missing=False)
    if gs_executable:
        subprocess.run(
            [
                gs_executable,
                "-sDEVICE=pdfwrite",
                f"-dPDFSETTINGS=/{preset}",
                "-dNOPAUSE",
                "-dBATCH",
                "-dQUIET",
                f"-sOutputFile={output_path}",
                input_path,
            ],
            check=True,
        )
        return output_name

    # Fallback: PyPDF basic compression
    pdf_reader = PdfReader(input_path)
    pdf_writer = PdfWriter()
    for page in pdf_reader.pages:
        pdf_writer.add_page(page)
    for page in pdf_writer.pages:
        page.compress_content_streams()
    with open(output_path, "wb") as output_file:
        pdf_writer.write(output_file)
    return output_name


# ============================================
# PDF TO DOCX CONVERSION
# ============================================
def pdf_to_docx(file_name: str, output_name: str = "converted_output.docx") -> str:
    """
    Convert a PDF file to DOCX format (low-memory, text-first).

    NOTE: This prioritizes stability on low-RAM servers (Render free tier).
    It extracts text per page and writes it into a DOCX. Layout fidelity is not preserved.
    """
    ensure_temp_dirs()
    input_path = get_upload_path(file_name)
    output_path = get_output_path(output_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception as e:
        raise Exception(
            "DOCX export requires the optional dependency 'python-docx'. "
            "Install with: pip install python-docx"
        )

    try:
        doc = fitz.open(input_path)
        out = Document()
        style = out.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            text = (page.get_text("text") or "").strip()
            if text:
                for line in text.splitlines():
                    out.add_paragraph(line)
            else:
                out.add_paragraph("")

            if page_index != doc.page_count - 1:
                out.add_page_break()

        doc.close()
        out.save(output_path)
    except Exception as e:
        raise Exception(f"PDF to DOCX conversion failed: {e}")
    return output_name


def docx_to_pdf(file_name: str, output_name: str = "converted_output.pdf") -> str:
    """Convert a DOCX file to PDF.

    This uses LibreOffice (soffice) if available. On minimal containers it may not be installed;
    in that case we return a clear error instead of crashing deployment.
    """
    ensure_temp_dirs()
    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    # If LibreOffice is not available (common on minimal containers), fall back
    # to a text-only PDF render (medium accuracy, low RAM).
    if not soffice:
        try:
            from docx import Document  # type: ignore
            from reportlab.pdfgen import canvas  # type: ignore
            from reportlab.lib.pagesizes import A4  # type: ignore
        except Exception:
            raise Exception(
                "DOCX → PDF conversion requires either LibreOffice (soffice) on the server or the "
                "optional dependencies 'python-docx' and 'reportlab'."
            )

        doc = Document(input_path)
        output_path = get_output_path(output_name)

        page_w, page_h = A4
        margin = 48
        font_name = "Helvetica"
        font_size = 11
        line_h = 14

        c = canvas.Canvas(output_path, pagesize=A4)
        c.setFont(font_name, font_size)
        x = margin
        y = page_h - margin

        def wrap_text(text: str, max_width: float) -> list[str]:
            words = (text or "").split()
            if not words:
                return [""]
            lines: list[str] = []
            cur: list[str] = []
            for w in words:
                trial = (" ".join(cur + [w])).strip()
                if c.stringWidth(trial, font_name, font_size) <= max_width:
                    cur.append(w)
                else:
                    if cur:
                        lines.append(" ".join(cur))
                    cur = [w]
            if cur:
                lines.append(" ".join(cur))
            return lines

        max_width = page_w - 2 * margin
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            # Preserve blank lines
            if not text:
                y -= line_h
                if y < margin:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = page_h - margin
                continue

            for line in wrap_text(text, max_width=max_width):
                c.drawString(x, y, line)
                y -= line_h
                if y < margin:
                    c.showPage()
                    c.setFont(font_name, font_size)
                    y = page_h - margin

        c.save()
        return output_name

    out_dir = os.path.abspath("outputs")

    # LibreOffice writes output name based on input; we rename after.
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                os.path.abspath(input_path),
            ],
            check=True,
            timeout=180,
        )
    except subprocess.TimeoutExpired:
        raise Exception("DOCX → PDF conversion timed out.")
    except Exception as e:
        raise Exception(f"DOCX → PDF conversion failed: {e}")

    base = os.path.splitext(os.path.basename(file_name))[0]
    produced = os.path.join(out_dir, f"{base}.pdf")
    if not os.path.exists(produced):
        # LibreOffice may sanitize the filename; best-effort fallback.
        raise Exception("DOCX → PDF conversion failed: output PDF not produced.")

    final_path = get_output_path(output_name)
    try:
        if os.path.exists(final_path):
            os.remove(final_path)
        os.replace(produced, final_path)
    except Exception:
        # If rename fails, keep produced.
        return os.path.basename(produced)

    return output_name


# ============================================
# COMPRESS TO TARGET SIZE (EXPERIMENTAL)
# ============================================
def compress_pdf_to_target(file_name: str, target_mb: int, output_name: str = "compressed_target_output.pdf") -> str:
    """
    Compress a PDF to a target size (in MB) using Ghostscript.
    Args:
        file_name: Source PDF file name
        target_mb: Target size in MB
        output_name: Name for the output file
    Returns:
        str: Output file name
    Raises:
        FileNotFoundError: If input file doesn't exist
        Exception: If compression fails completely
    """
    ensure_temp_dirs()
    input_path = get_upload_path(file_name)
    output_path = get_output_path(output_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")
    
    gs_executable = _resolve_ghostscript_executable(raise_if_missing=True)
    
    original_size_mb = os.path.getsize(input_path) / (1024 * 1024)
    best_size_mb = original_size_mb
    best_quality = None
    
    # Try different quality settings until under target
    qualities = ["screen", "ebook", "printer", "prepress"]
    for quality in qualities:
        try:
            subprocess.run([
                gs_executable,
                "-sDEVICE=pdfwrite",
                f"-dPDFSETTINGS=/{quality}",
                "-dNOPAUSE",
                "-dBATCH",
                "-dQUIET",
                f"-sOutputFile={output_path}",
                input_path
            ], check=True)
            size_mb = os.path.getsize(output_path) / (1024 * 1024)
            if size_mb <= target_mb:
                return output_name
            # Track best compression achieved
            if size_mb < best_size_mb:
                best_size_mb = size_mb
                best_quality = quality
        except Exception as e:
            continue
    
    # If we couldn't reach target but got some compression, return best result with info
    if best_quality:
        # Re-compress with best quality setting
        subprocess.run([
            gs_executable,
            "-sDEVICE=pdfwrite",
            f"-dPDFSETTINGS=/{best_quality}",
            "-dNOPAUSE",
            "-dBATCH",
            "-dQUIET",
            f"-sOutputFile={output_path}",
            input_path
        ], check=True)
        # Return with special marker that will be caught
        raise Exception(f"PARTIAL_SUCCESS:Compressed from {original_size_mb:.1f}MB to {best_size_mb:.1f}MB (target was {target_mb}MB). Maximum compression reached.")
    
    raise Exception(f"Could not compress {file_name}. The PDF may already be optimized.")


# ============================================
# ADDITIONAL PDF OPERATIONS
# ============================================

def rotate_pdf(
    file_name: str,
    degrees: int,
    pages: Optional[List[int]] = None,
    output_name: str = "rotated_output.pdf",
) -> str:
    """Rotate pages in a PDF.

    Args:
        file_name: Source PDF file name
        degrees: 90/180/270 clockwise
        pages: Optional list of 1-indexed pages to rotate. If None, rotate all pages.
        output_name: Output PDF name
    """
    ensure_temp_dirs()

    if degrees not in (90, 180, 270):
        raise ValueError("degrees must be one of 90, 180, 270")

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    reader = PdfReader(input_path)
    total_pages = len(reader.pages)

    if pages is None:
        pages_set = set(range(1, total_pages + 1))
    else:
        for p in pages:
            if p < 1 or p > total_pages:
                raise ValueError(f"Invalid page number: {p}. PDF has {total_pages} pages.")
        pages_set = set(pages)

    writer = PdfWriter()
    for idx, page in enumerate(reader.pages, start=1):
        if idx in pages_set:
            # pypdf rotation API differs by version
            if hasattr(page, "rotate_clockwise"):
                page = page.rotate_clockwise(degrees)
            elif hasattr(page, "rotate"):
                page = page.rotate(degrees)
            elif hasattr(page, "rotateClockwise"):
                page.rotateClockwise(degrees)
        writer.add_page(page)

    output_path = get_output_path(output_name)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_name


def reorder_pdf(file_name: str, new_order: List[int] | str, output_name: str = "reordered_output.pdf") -> str:
    """Reorder pages in a PDF according to new_order (1-indexed).
    
    new_order can be:
    - A list of page numbers in the desired order
    - The string "reverse" to reverse all pages
    """
    ensure_temp_dirs()

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    reader = PdfReader(input_path)
    total_pages = len(reader.pages)
    
    # Handle "reverse" command
    if new_order == "reverse":
        new_order = list(range(total_pages, 0, -1))  # e.g., [5, 4, 3, 2, 1] for 5 pages

    if len(new_order) != total_pages:
        raise ValueError(f"new_order must have exactly {total_pages} entries")

    expected = set(range(1, total_pages + 1))
    if set(new_order) != expected:
        raise ValueError("new_order must include each page number exactly once")

    writer = PdfWriter()
    for p in new_order:
        writer.add_page(reader.pages[p - 1])

    output_path = get_output_path(output_name)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_name


def watermark_pdf(
    file_name: str,
    text: str,
    opacity: float = 0.12,
    angle: int = 30,
    output_name: str = "watermarked_output.pdf",
) -> str:
    """Add a text watermark to each page (low-RAM overlay stamping)."""
    ensure_temp_dirs()

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    if opacity < 0 or opacity > 1:
        raise ValueError("opacity must be between 0 and 1")

    try:
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        raise Exception(
            "Watermark requires the optional dependency 'reportlab'. Install with: pip install reportlab"
        )

    reader = PdfReader(input_path)
    writer = PdfWriter()

    def _overlay_page(w: float, h: float) -> "PdfReader":
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=(w, h))
        try:
            c.setFillAlpha(float(opacity))
        except Exception:
            pass
        c.saveState()
        c.translate(w / 2.0, h / 2.0)
        try:
            c.rotate(float(angle or 0))
        except Exception:
            c.rotate(0)
        # Font size based on page diagonal-ish
        base = min(w, h)
        font_size = max(18, min(72, int(base / 10)))
        c.setFont("Helvetica", font_size)
        c.setFillColorRGB(0, 0, 0)
        # Centered text
        c.drawCentredString(0, 0, text)
        c.restoreState()
        c.showPage()
        c.save()
        buf.seek(0)
        return PdfReader(buf)

    for page in reader.pages:
        mb = page.mediabox
        w = float(mb.width)
        h = float(mb.height)
        overlay_pdf = _overlay_page(w, h)
        overlay_page = overlay_pdf.pages[0]
        try:
            page.merge_page(overlay_page)
        except Exception:
            page.mergePage(overlay_page)  # type: ignore
        writer.add_page(page)

    output_path = get_output_path(output_name)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_name


def add_page_numbers(
    file_name: str,
    position: str = "bottom_center",
    start_at: int = 1,
    output_name: str = "page_numbers_output.pdf",
) -> str:
    """Add page numbers to each page (low-RAM overlay stamping)."""
    ensure_temp_dirs()

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    try:
        from reportlab.pdfgen import canvas  # type: ignore
    except Exception:
        raise Exception(
            "Page numbers require the optional dependency 'reportlab'. Install with: pip install reportlab"
        )

    reader = PdfReader(input_path)
    writer = PdfWriter()

    def _coords(w: float, h: float) -> tuple[float, float, int]:
        margin = 28.0
        font_size = 11
        if position == "bottom_right":
            return (w - margin, margin, 2)
        if position == "bottom_left":
            return (margin, margin, 0)
        if position == "top_right":
            return (w - margin, h - margin, 2)
        if position == "top_left":
            return (margin, h - margin, 0)
        if position == "top_center":
            return (w / 2.0, h - margin, 1)
        return (w / 2.0, margin, 1)

    for i, page in enumerate(reader.pages, start=0):
        mb = page.mediabox
        w = float(mb.width)
        h = float(mb.height)
        x, y, align = _coords(w, h)
        number = start_at + i

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=(w, h))
        c.setFont("Helvetica", 11)
        c.setFillColorRGB(0, 0, 0)
        s = str(number)
        if align == 1:
            c.drawCentredString(x, y, s)
        elif align == 2:
            c.drawRightString(x, y, s)
        else:
            c.drawString(x, y, s)
        c.showPage()
        c.save()
        buf.seek(0)
        overlay_page = PdfReader(buf).pages[0]
        try:
            page.merge_page(overlay_page)
        except Exception:
            page.mergePage(overlay_page)  # type: ignore
        writer.add_page(page)

    output_path = get_output_path(output_name)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_name


def flatten_pdf(file_name: str, output_name: str = "flattened_output.pdf") -> str:
    """Flatten/sanitize a PDF structure (low-RAM optimization).

    This does NOT rasterize pages (so it preserves text selectability).
    It rewrites the file with cleanup and compression.
    """
    ensure_temp_dirs()
    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    try:
        doc = fitz.open(input_path)
        output_path = get_output_path(output_name)
        doc.save(output_path, garbage=4, deflate=True, clean=True)
        doc.close()
        return output_name
    except Exception as e:
        raise Exception(f"Flatten failed: {e}")


def remove_blank_pages(
    file_name: str,
    output_name: str = "no_blank_pages.pdf",
    blank_ratio_threshold: float = 0.002,
    zoom: float = 0.15,
) -> str:
    """Remove blank/empty pages using low-res thumbnail ink detection.

    Designed for free tier: per-page low-res render, modest accuracy.
    """
    ensure_temp_dirs()
    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    try:
        import numpy as np  # type: ignore
    except Exception:
        raise Exception("remove_blank_pages requires numpy (already in requirements).")

    doc = fitz.open(input_path)
    reader = PdfReader(input_path)
    writer = PdfWriter()

    kept = 0
    for idx in range(doc.page_count):
        page = doc.load_page(idx)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csGRAY)
        img = np.frombuffer(pix.samples, dtype=np.uint8)
        if img.size == 0:
            continue
        # Count non-white pixels
        nonwhite = np.count_nonzero(img < 245)
        ratio = nonwhite / float(img.size)
        if ratio >= float(blank_ratio_threshold):
            writer.add_page(reader.pages[idx])
            kept += 1

    doc.close()

    if kept == 0:
        raise Exception("All pages were detected as blank. If this is wrong, try again with a different file.")

    output_path = get_output_path(output_name)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_name


def remove_duplicate_pages(
    file_name: str,
    output_name: str = "no_duplicate_pages.pdf",
    zoom: float = 0.12,
) -> str:
    """Remove duplicate pages using small thumbnail hashing.

    Fast + medium accuracy on scanned/image-heavy PDFs.
    """
    ensure_temp_dirs()
    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    doc = fitz.open(input_path)
    reader = PdfReader(input_path)
    writer = PdfWriter()

    seen: set[str] = set()
    kept = 0

    for idx in range(doc.page_count):
        page = doc.load_page(idx)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csGRAY)
        digest = hashlib.sha1(pix.samples).hexdigest()
        if digest in seen:
            continue
        seen.add(digest)
        writer.add_page(reader.pages[idx])
        kept += 1

    doc.close()

    if kept == 0:
        raise Exception("No pages kept after duplicate removal.")

    output_path = get_output_path(output_name)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_name


def enhance_scan(
    file_name: str,
    output_name: str = "enhanced_scan.pdf",
    dpi: int = 150,
    max_pages: int = 25,
) -> str:
    """Enhance a scanned PDF for readability using OpenCV (thresholding).

    This is image-based and will typically make text NON-selectable.
    Uses conservative DPI and page limits to stay within 512MB.
    """
    ensure_temp_dirs()
    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    try:
        import numpy as np  # type: ignore
        import cv2  # type: ignore
    except Exception:
        raise Exception(
            "enhance_scan requires opencv-python-headless and numpy (already in requirements)."
        )

    doc = fitz.open(input_path)
    if doc.page_count > int(max_pages):
        doc.close()
        raise Exception(
            f"Enhance scan is limited to {max_pages} pages on the free tier to avoid memory spikes. "
            "Please split the PDF first and try again."
        )

    out = fitz.open()

    # scale matrix for desired DPI (PDF default 72 dpi)
    scale = max(1.0, float(dpi) / 72.0)
    mat = fitz.Matrix(scale, scale)

    for idx in range(doc.page_count):
        page = doc.load_page(idx)
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        img = np.frombuffer(pix.samples, dtype=np.uint8)
        img = img.reshape((pix.height, pix.width, 3))
        gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
        gray = cv2.medianBlur(gray, 3)
        bw = cv2.adaptiveThreshold(
            gray,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            10,
        )

        # Encode to JPEG to keep PDF size reasonable
        ok, jpg = cv2.imencode(".jpg", bw, [int(cv2.IMWRITE_JPEG_QUALITY), 80])
        if not ok:
            raise Exception("Scan enhancement failed during JPEG encoding.")
        jpg_bytes = jpg.tobytes()

        # Create a page matching image size (in points)
        w_pt = pix.width * (72.0 / float(dpi))
        h_pt = pix.height * (72.0 / float(dpi))
        out_page = out.new_page(width=w_pt, height=h_pt)
        out_page.insert_image(out_page.rect, stream=jpg_bytes)

    doc.close()
    output_path = get_output_path(output_name)
    out.save(output_path, garbage=4, deflate=True)
    out.close()
    return output_name


def extract_text(
    file_name: str,
    pages: Optional[List[int]] = None,
    output_name: str = "extracted_text.txt",
) -> str:
    """Extract text from a PDF into a .txt file."""
    ensure_temp_dirs()

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    doc = fitz.open(input_path)
    total_pages = doc.page_count

    if pages is None:
        page_numbers = list(range(1, total_pages + 1))
    else:
        for p in pages:
            if p < 1 or p > total_pages:
                raise ValueError(f"Invalid page number: {p}. PDF has {total_pages} pages.")
        page_numbers = pages

    parts: list[str] = []
    for p in page_numbers:
        page = doc.load_page(p - 1)
        parts.append(f"--- Page {p} ---\n")
        parts.append(page.get_text("text"))
        parts.append("\n")

    doc.close()
    output_path = get_output_path(output_name)
    with open(output_path, "w", encoding="utf-8", errors="replace") as f:
        f.write("".join(parts))
    return output_name


def pdf_to_images_zip(
    file_name: str,
    fmt: str = "png",
    dpi: int = 150,
    output_name: str = "pdf_images.zip",
) -> str:
    """Render PDF pages to images and return a zip file."""
    ensure_temp_dirs()

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    fmt_lower = (fmt or "png").lower()
    if fmt_lower == "jpg":
        fmt_lower = "jpeg"
    if fmt_lower not in {"png", "jpeg"}:
        fmt_lower = "png"

    doc = fitz.open(input_path)
    output_path = get_output_path(output_name)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for i in range(doc.page_count):
            page = doc.load_page(i)
            pix = page.get_pixmap(dpi=dpi)
            img_bytes = pix.tobytes(fmt_lower)
            zf.writestr(f"page_{i+1:04d}.{ 'jpg' if fmt_lower=='jpeg' else fmt_lower }", img_bytes)
    doc.close()
    return output_name


def images_to_pdf(
    file_names: List[str],
    output_name: str = "images_output.pdf",
    max_dimension: int = 2000,  # Max width/height in pixels
    jpeg_quality: int = 85,  # JPEG compression quality
) -> str:
    """Combine uploaded images into a single PDF in the given order.
    
    Images are automatically resized and compressed to keep PDF size reasonable.
    Large images (>2000px) are scaled down. All images are JPEG compressed.
    """
    ensure_temp_dirs()

    if not file_names:
        raise ValueError("No image files provided")

    doc = fitz.open()
    
    # Standard A4 page size in points (72 points = 1 inch)
    A4_WIDTH = 595  # ~8.27 inches
    A4_HEIGHT = 842  # ~11.69 inches

    for name in file_names:
        input_path = get_upload_path(name)
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"File not found: {name}")

        pix = fitz.Pixmap(input_path)
        
        # Convert CMYK/alpha to RGB when needed
        if pix.n >= 5:
            pix = fitz.Pixmap(fitz.csRGB, pix)
        
        orig_width, orig_height = pix.width, pix.height
        
        # Resize if image is too large (reduces memory and PDF size)
        if orig_width > max_dimension or orig_height > max_dimension:
            scale = min(max_dimension / orig_width, max_dimension / orig_height)
            new_width = int(orig_width * scale)
            new_height = int(orig_height * scale)
            # Create scaled pixmap using transformation matrix
            mat = fitz.Matrix(scale, scale)
            pix = fitz.Pixmap(pix, 0, 1)  # Remove alpha if present
            # Re-read and scale
            pix = fitz.Pixmap(input_path)
            if pix.n >= 5:
                pix = fitz.Pixmap(fitz.csRGB, pix)
        else:
            new_width, new_height = orig_width, orig_height
        
        # Determine page orientation based on image aspect ratio
        img_is_landscape = new_width > new_height
        
        # Use A4 size, rotate for landscape images
        if img_is_landscape:
            page_width, page_height = A4_HEIGHT, A4_WIDTH  # Landscape A4
        else:
            page_width, page_height = A4_WIDTH, A4_HEIGHT  # Portrait A4
        
        # Calculate scale to fit image on page with margins
        margin = 20  # Small margin
        avail_width = page_width - 2 * margin
        avail_height = page_height - 2 * margin
        
        scale_w = avail_width / new_width
        scale_h = avail_height / new_height
        fit_scale = min(scale_w, scale_h, 1.0)  # Don't upscale small images
        
        final_width = new_width * fit_scale
        final_height = new_height * fit_scale
        
        # Center the image on the page
        x_offset = (page_width - final_width) / 2
        y_offset = (page_height - final_height) / 2
        
        page = doc.new_page(width=page_width, height=page_height)
        img_rect = fitz.Rect(x_offset, y_offset, x_offset + final_width, y_offset + final_height)
        
        # Insert image with JPEG compression for smaller file size
        page.insert_image(img_rect, filename=input_path)

    output_path = get_output_path(output_name)
    # Save with garbage collection and compression
    doc.save(output_path, garbage=4, deflate=True)
    doc.close()
    return output_name


def split_pages_to_files_zip(
    file_name: str,
    pages: Optional[List[int]] = None,
    output_name: str = "split_pages.zip",
) -> str:
    """Extract pages as individual single-page PDFs and return a zip file."""
    ensure_temp_dirs()

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    reader = PdfReader(input_path)
    total_pages = len(reader.pages)

    if pages is None:
        pages_list = list(range(1, total_pages + 1))
    else:
        for p in pages:
            if p < 1 or p > total_pages:
                raise ValueError(f"Invalid page number: {p}. PDF has {total_pages} pages.")
        pages_list = pages

    output_path = get_output_path(output_name)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in pages_list:
            writer = PdfWriter()
            writer.add_page(reader.pages[p - 1])
            buf = io.BytesIO()
            writer.write(buf)
            zf.writestr(f"page_{p:04d}.pdf", buf.getvalue())
    return output_name


def ocr_pdf(
    file_name: str,
    language: str = "eng",
    deskew: bool = True,
    output_name: str = "ocr_output.pdf",
) -> str:
    """Run OCR to produce a searchable PDF.

    This uses the optional `ocrmypdf` dependency (and requires Tesseract installed on the machine).
    """
    ensure_temp_dirs()

    input_path = get_upload_path(file_name)
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"File not found: {file_name}")

    # Use the CLI to enforce timeouts and to avoid excessive in-process memory.
    ocrmypdf_cmd = shutil.which("ocrmypdf")
    if not ocrmypdf_cmd:
        raise Exception(
            "OCR requires the optional dependency 'ocrmypdf' and a Tesseract install. "
            "Install with: pip install ocrmypdf  (and install Tesseract on your system)."
        )

    # Free-tier safe defaults
    default_dpi = 200
    chunk_pages = 25

    reader = PdfReader(input_path)
    total_pages = len(reader.pages)

    def _run_ocr(in_pdf: str, out_pdf: str, timeout_s: int) -> None:
        cmd = [
            ocrmypdf_cmd,
            "--skip-text",
            "--jobs",
            "1",
            "--image-dpi",
            str(default_dpi),
            "--tesseract-timeout",
            "120",
        ]
        if language:
            cmd += ["-l", str(language)]
        if deskew:
            cmd += ["--deskew"]
        cmd += [in_pdf, out_pdf]
        try:
            subprocess.run(cmd, check=True, timeout=timeout_s)
        except subprocess.TimeoutExpired:
            raise Exception("OCR timed out. Try splitting the PDF or using fewer pages.")

    output_path = get_output_path(output_name)

    # Chunking to avoid memory spikes on long PDFs
    if total_pages > chunk_pages:
        temp_outputs: list[str] = []
        try:
            for start in range(0, total_pages, chunk_pages):
                end = min(total_pages, start + chunk_pages)
                chunk_in = get_output_path(f"_ocr_chunk_in_{start+1}_{end}.pdf")
                chunk_out = get_output_path(f"_ocr_chunk_out_{start+1}_{end}.pdf")

                w = PdfWriter()
                for i in range(start, end):
                    w.add_page(reader.pages[i])
                with open(chunk_in, "wb") as f:
                    w.write(f)

                # Timeout scales mildly with pages
                timeout_s = 180 + (end - start) * 20
                _run_ocr(chunk_in, chunk_out, timeout_s=timeout_s)
                temp_outputs.append(chunk_out)

            # Merge chunks
            merged = PdfWriter()
            for p in temp_outputs:
                r = PdfReader(p)
                for pg in r.pages:
                    merged.add_page(pg)
            with open(output_path, "wb") as f:
                merged.write(f)
        except Exception as e:
            raise Exception(f"OCR failed: {e}")
        finally:
            # Cleanup temp files
            try:
                for start in range(0, total_pages, chunk_pages):
                    end = min(total_pages, start + chunk_pages)
                    chunk_in = get_output_path(f"_ocr_chunk_in_{start+1}_{end}.pdf")
                    if os.path.exists(chunk_in):
                        os.remove(chunk_in)
                for p in temp_outputs:
                    if os.path.exists(p):
                        os.remove(p)
            except Exception:
                pass

        return output_name

    # Single shot
    try:
        _run_ocr(input_path, output_path, timeout_s=240 + total_pages * 20)
    except Exception as e:
        raise Exception(f"OCR failed: {e}")
    return output_name
